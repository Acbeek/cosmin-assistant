"""Provisional DOCX export interface and stub implementation."""

from __future__ import annotations

from pathlib import Path
from typing import Protocol

from docx import Document


class DocxExporter(Protocol):
    """Interface for DOCX exporters used by table/output builders."""

    def export_summary(self, *, output_path: Path, report_markdown: str) -> Path:
        """Export a summary report to DOCX."""


class ProvisionalDocxExporter:
    """Stub DOCX exporter for Task 10 provisional end-to-end outputs."""

    def export_summary(self, *, output_path: Path, report_markdown: str) -> Path:
        document = Document()
        document.add_heading("COSMIN Assistant Provisional Export", level=1)
        document.add_paragraph(
            "This DOCX file is a provisional stub. Final COSMIN table templates will be "
            "implemented in a later task."
        )
        document.add_heading("Summary Markdown Snapshot", level=2)
        document.add_paragraph(report_markdown)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path
