"""DOCX exports for COSMIN-style Template 5/7/8 intermediate tables.

This module consumes intermediate table objects and remains independent from
scoring/synthesis logic.
"""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from typing import Generic, TypeVar

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell, _Row

from cosmin_assistant.tables.intermediate_models import (
    TableLegendEntry,
    Template5CharacteristicsRow,
    Template5CharacteristicsTable,
    Template7EvidenceRow,
    Template7EvidenceTable,
    Template7RowKind,
    Template8SummaryRow,
    Template8SummaryTable,
)

_HEADER_SHADE = "EDEDED"
_ALTERNATING_SHADE = "F7FBFF"
_FONT_SIZE_PT = 9

_RowT = TypeVar(
    "_RowT",
    "_Template5Payload",
    "_Template7Payload",
    "_Template8Payload",
)


@dataclass(frozen=True)
class _Template5Payload:
    row: Template5CharacteristicsRow
    blank_group_fields: bool


@dataclass(frozen=True)
class _Template7Payload:
    row: Template7EvidenceRow
    blank_group_fields: bool
    summary_label: bool


@dataclass(frozen=True)
class _Template8Payload:
    row: Template8SummaryRow
    blank_group_fields: bool


@dataclass(frozen=True)
class _ColumnSpec(Generic[_RowT]):
    header: str
    value_getter: Callable[[_RowT], str]
    align_right: bool = False


class CosminTableDocxExporter:
    """DOCX exporter for COSMIN-style Template 5/7/8 tables."""

    def export_template5(
        self,
        *,
        table: Template5CharacteristicsTable,
        output_path: str | Path,
    ) -> Path:
        """Export Template 5 equivalent table to DOCX."""

        document = Document()
        document.add_heading(table.title, level=1)
        specs = _template5_columns()
        rows = tuple(
            _Template5Payload(
                row=row,
                blank_group_fields=row.is_additional_study_row,
            )
            for row in table.rows
        )
        _render_table(document=document, rows=rows, columns=specs, group_key_fn=_group_key_t5)
        _append_legends(document, table.legends)
        return _save_document(document, output_path)

    def export_template7(
        self,
        *,
        table: Template7EvidenceTable,
        output_path: str | Path,
    ) -> Path:
        """Export Template 7 equivalent table to DOCX."""

        document = Document()
        document.add_heading(table.title, level=1)
        specs = _template7_columns()

        rows: list[_Template7Payload] = []
        previous_group_key: tuple[str, str | None, str | None, str] | None = None
        for row in table.rows:
            group_key = (
                row.instrument_name,
                row.instrument_version,
                row.subscale,
                row.measurement_property,
            )
            if row.row_kind is Template7RowKind.STUDY:
                blank_group_fields = bool(row.is_additional_study_row)
                rows.append(
                    _Template7Payload(
                        row=row,
                        blank_group_fields=blank_group_fields,
                        summary_label=False,
                    )
                )
            else:
                blank_group_fields = previous_group_key == group_key
                rows.append(
                    _Template7Payload(
                        row=row,
                        blank_group_fields=blank_group_fields,
                        summary_label=True,
                    )
                )
            previous_group_key = group_key

        _render_table(
            document=document,
            rows=tuple(rows),
            columns=specs,
            group_key_fn=_group_key_t7,
        )
        _append_legends(document, table.legends)
        return _save_document(document, output_path)

    def export_template8(
        self,
        *,
        table: Template8SummaryTable,
        output_path: str | Path,
    ) -> Path:
        """Export Template 8 equivalent table to DOCX."""

        document = Document()
        document.add_heading(table.title, level=1)
        specs = _template8_columns()

        rows: list[_Template8Payload] = []
        previous_group_key: tuple[str, str | None, str | None] | None = None
        for row in table.rows:
            group_key = (row.instrument_name, row.instrument_version, row.subscale)
            rows.append(
                _Template8Payload(
                    row=row,
                    blank_group_fields=previous_group_key == group_key,
                )
            )
            previous_group_key = group_key

        _render_table(
            document=document,
            rows=tuple(rows),
            columns=specs,
            group_key_fn=_group_key_t8,
        )
        _append_legends(document, table.legends)
        return _save_document(document, output_path)


def export_template5_docx(
    *,
    table: Template5CharacteristicsTable,
    output_path: str | Path,
) -> Path:
    """Export Template 5 equivalent table to DOCX."""

    return CosminTableDocxExporter().export_template5(table=table, output_path=output_path)


def export_template7_docx(
    *,
    table: Template7EvidenceTable,
    output_path: str | Path,
) -> Path:
    """Export Template 7 equivalent table to DOCX."""

    return CosminTableDocxExporter().export_template7(table=table, output_path=output_path)


def export_template8_docx(
    *,
    table: Template8SummaryTable,
    output_path: str | Path,
) -> Path:
    """Export Template 8 equivalent table to DOCX."""

    return CosminTableDocxExporter().export_template8(table=table, output_path=output_path)


def _render_table(
    *,
    document: DocumentObject,
    rows: tuple[_RowT, ...],
    columns: tuple[_ColumnSpec[_RowT], ...],
    group_key_fn: Callable[[_RowT], tuple[str, str | None, str | None]],
) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    _write_header_row(row=table.rows[0], columns=columns)

    group_colors: dict[tuple[str, str | None, str | None], str | None] = {}
    next_group_index = 0
    for row_payload in rows:
        group_key = group_key_fn(row_payload)
        if group_key not in group_colors:
            group_colors[group_key] = _ALTERNATING_SHADE if next_group_index % 2 == 1 else None
            next_group_index += 1
        row_color = group_colors[group_key]

        row = table.add_row()
        _write_data_row(row=row, payload=row_payload, columns=columns, row_shading=row_color)


def _write_header_row(*, row: _Row, columns: tuple[_ColumnSpec[_RowT], ...]) -> None:
    for index, spec in enumerate(columns):
        cell = row.cells[index]
        _set_cell_text(
            cell=cell,
            text=spec.header,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            bold=True,
            shading=_HEADER_SHADE,
        )
    _set_repeat_table_header(row)


def _write_data_row(
    *,
    row: _Row,
    payload: _RowT,
    columns: tuple[_ColumnSpec[_RowT], ...],
    row_shading: str | None,
) -> None:
    for index, spec in enumerate(columns):
        text = spec.value_getter(payload)
        alignment = WD_ALIGN_PARAGRAPH.RIGHT if spec.align_right else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_text(
            cell=row.cells[index],
            text=text,
            alignment=alignment,
            bold=False,
            shading=row_shading,
        )


def _set_cell_text(
    *,
    cell: _Cell,
    text: str,
    alignment: WD_ALIGN_PARAGRAPH,
    bold: bool,
    shading: str | None,
) -> None:
    cell.text = text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    for run in paragraph.runs:
        run.bold = bold
        run.font.size = Pt(_FONT_SIZE_PT)
    if shading is not None:
        _apply_cell_shading(cell=cell, fill=shading)


def _apply_cell_shading(*, cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_repeat_table_header(row: _Row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _append_legends(document: DocumentObject, legends: tuple[TableLegendEntry, ...]) -> None:
    if not legends:
        return

    heading = document.add_paragraph()
    run = heading.add_run("Legend")
    run.bold = True
    run.font.size = Pt(_FONT_SIZE_PT)

    for legend in legends:
        paragraph = document.add_paragraph(f"{legend.key}: {legend.description}")
        for run in paragraph.runs:
            run.font.size = Pt(_FONT_SIZE_PT)


def _save_document(document: DocumentObject, output_path: str | Path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


def _template5_columns() -> tuple[_ColumnSpec[_Template5Payload], ...]:
    return (
        _ColumnSpec(
            header="PROM",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else payload.row.instrument_name
            ),
        ),
        _ColumnSpec(
            header="Version",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else _stringify(payload.row.instrument_version)
            ),
        ),
        _ColumnSpec(
            header="Subscale",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else _stringify(payload.row.subscale)
            ),
        ),
        _ColumnSpec(header="Study", value_getter=lambda payload: payload.row.study_id),
        _ColumnSpec(
            header="Study Design",
            value_getter=lambda payload: _stringify(payload.row.study_design),
        ),
        _ColumnSpec(
            header="Target Population",
            value_getter=lambda payload: _stringify(payload.row.target_population),
        ),
        _ColumnSpec(
            header="Language",
            value_getter=lambda payload: _stringify(payload.row.language),
        ),
        _ColumnSpec(
            header="Country",
            value_getter=lambda payload: _stringify(payload.row.country),
        ),
        _ColumnSpec(
            header="Enrollment n",
            value_getter=lambda payload: _number(payload.row.enrollment_n),
            align_right=True,
        ),
        _ColumnSpec(
            header="Analyzed n",
            value_getter=lambda payload: _number(payload.row.analyzed_n),
            align_right=True,
        ),
        _ColumnSpec(
            header="Limb-level n",
            value_getter=lambda payload: _number(payload.row.limb_level_n),
            align_right=True,
        ),
        _ColumnSpec(
            header="Follow-up",
            value_getter=lambda payload: _stringify(payload.row.follow_up_schedule),
        ),
        _ColumnSpec(
            header="Measurement Properties Mentioned",
            value_getter=lambda payload: _stringify(payload.row.measurement_properties_mentioned),
        ),
    )


def _template7_columns() -> tuple[_ColumnSpec[_Template7Payload], ...]:
    return (
        _ColumnSpec(
            header="PROM",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else payload.row.instrument_name
            ),
        ),
        _ColumnSpec(
            header="Version",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else _stringify(payload.row.instrument_version)
            ),
        ),
        _ColumnSpec(
            header="Subscale",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else _stringify(payload.row.subscale)
            ),
        ),
        _ColumnSpec(
            header="Measurement Property",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else payload.row.measurement_property
            ),
        ),
        _ColumnSpec(
            header="Study",
            value_getter=lambda payload: (
                "Summary" if payload.summary_label else _stringify(payload.row.study_id)
            ),
        ),
        _ColumnSpec(
            header="RoB",
            value_getter=lambda payload: _stringify(payload.row.per_study_rob),
        ),
        _ColumnSpec(
            header="Per-study Result",
            value_getter=lambda payload: _stringify(payload.row.per_study_result),
        ),
        _ColumnSpec(
            header="Study Rating",
            value_getter=lambda payload: _stringify(payload.row.study_rating),
        ),
        _ColumnSpec(
            header="Summarized Result",
            value_getter=lambda payload: _stringify(payload.row.summarized_result),
        ),
        _ColumnSpec(
            header="Overall Rating",
            value_getter=lambda payload: _stringify(payload.row.overall_rating),
        ),
        _ColumnSpec(
            header="Certainty of Evidence",
            value_getter=lambda payload: _stringify(payload.row.certainty_of_evidence),
        ),
        _ColumnSpec(
            header="Total n",
            value_getter=lambda payload: _number(payload.row.total_sample_size),
            align_right=True,
        ),
    )


def _template8_columns() -> tuple[_ColumnSpec[_Template8Payload], ...]:
    return (
        _ColumnSpec(
            header="PROM",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else payload.row.instrument_name
            ),
        ),
        _ColumnSpec(
            header="Version",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else _stringify(payload.row.instrument_version)
            ),
        ),
        _ColumnSpec(
            header="Subscale",
            value_getter=lambda payload: (
                "" if payload.blank_group_fields else _stringify(payload.row.subscale)
            ),
        ),
        _ColumnSpec(
            header="Measurement Property",
            value_getter=lambda payload: payload.row.measurement_property,
        ),
        _ColumnSpec(
            header="Summarized Result",
            value_getter=lambda payload: _stringify(payload.row.summarized_result),
        ),
        _ColumnSpec(
            header="Overall Rating",
            value_getter=lambda payload: _stringify(payload.row.overall_rating),
        ),
        _ColumnSpec(
            header="Certainty of Evidence",
            value_getter=lambda payload: _stringify(payload.row.certainty_of_evidence),
        ),
        _ColumnSpec(
            header="Total n",
            value_getter=lambda payload: _number(payload.row.total_sample_size),
            align_right=True,
        ),
        _ColumnSpec(
            header="Inconsistent Findings",
            value_getter=lambda payload: _bool_text(payload.row.inconsistent_findings),
        ),
    )


def _group_key_t5(payload: _Template5Payload) -> tuple[str, str | None, str | None]:
    row = payload.row
    return (row.instrument_name, row.instrument_version, row.subscale)


def _group_key_t7(payload: _Template7Payload) -> tuple[str, str | None, str | None]:
    row = payload.row
    return (row.instrument_name, row.instrument_version, row.subscale)


def _group_key_t8(payload: _Template8Payload) -> tuple[str, str | None, str | None]:
    row = payload.row
    return (row.instrument_name, row.instrument_version, row.subscale)


def _number(value: int | None) -> str:
    return "" if value is None else str(value)


def _bool_text(value: bool | None) -> str:
    if value is None:
        return ""
    return "yes" if value else "no"


def _stringify(value: str | None) -> str:
    return value or ""
