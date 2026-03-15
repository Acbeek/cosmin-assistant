"""Structural tests for Template 5/6/7/8 DOCX table exports."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import _Cell, _Row

from cosmin_assistant.tables import (
    export_template5_docx,
    export_template6_docx,
    export_template7_docx,
    export_template8_docx,
)
from cosmin_assistant.tables.intermediate_models import (
    TableLegendEntry,
    Template5CharacteristicsRow,
    Template5CharacteristicsTable,
    Template6ContentValidityRow,
    Template6ContentValidityTable,
    Template6RowKind,
    Template7EvidenceRow,
    Template7EvidenceTable,
    Template7RowKind,
    Template8SummaryRow,
    Template8SummaryTable,
)


def test_template5_docx_structure_grouping_and_legend(tmp_path: Path) -> None:
    table = Template5CharacteristicsTable(
        id="table.5",
        rows=(
            Template5CharacteristicsRow(
                id="t5row.1",
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                study_id="study.1",
                study_display_label="Hafner et al., 2022",
                study_order_within_instrument=1,
                is_additional_study_row=False,
                study_design="prospective",
                target_population="tfa",
                language="english",
                country="us",
                enrollment_n=41,
                analyzed_n=37,
                limb_level_n=51,
                follow_up_schedule="baseline, 24 months",
                measurement_properties_mentioned="reliability",
            ),
            Template5CharacteristicsRow(
                id="t5row.2",
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                study_id="study.2",
                study_display_label="Hafner et al., 2022",
                study_order_within_instrument=2,
                is_additional_study_row=True,
                study_design="prospective",
                target_population="tfa",
            ),
            Template5CharacteristicsRow(
                id="t5row.3",
                instrument_name="PROM-X",
                instrument_version="v2",
                subscale="pain",
                study_id="study.3",
                study_display_label="Hafner et al., 2022",
                study_order_within_instrument=1,
                is_additional_study_row=False,
                study_design="cross-sectional",
                target_population="tfa",
                enrollment_n=88,
            ),
        ),
        legends=(
            TableLegendEntry(
                key="additional_study_row",
                description="Rows below are additional study reports.",
            ),
            TableLegendEntry(
                key="blank_or_na",
                description="Blank cells indicate NA or not assessed.",
            ),
        ),
    )

    output_path = tmp_path / "template5.docx"
    export_template5_docx(table=table, output_path=output_path)

    doc = Document(str(output_path))
    assert len(doc.tables) == 1
    word_table = doc.tables[0]

    headers = [cell.text for cell in word_table.rows[0].cells]
    assert headers == [
        "PROM",
        "Version",
        "Subscale",
        "Study",
        "Study Design",
        "Target Population",
        "Language",
        "Country",
        "Enrollment n",
        "Analyzed n",
        "Limb-level n",
        "Follow-up",
        "Measurement Properties Mentioned",
    ]
    assert _has_repeat_header(word_table.rows[0]) is True

    rows = [[cell.text for cell in row.cells] for row in word_table.rows[1:]]
    assert rows[1][:3] == ["", "", ""]
    assert rows[0][3] == "Hafner et al., 2022"
    assert rows[1][3] == "Hafner et al., 2022"
    assert rows[0][8] == "41"
    assert word_table.rows[1].cells[8].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    # Alternating group shading should be present for at least one group.
    assert any(_cell_fill(row.cells[0]) == "F7FBFF" for row in word_table.rows[1:])

    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Legend" in document_text
    assert "additional_study_row" in document_text
    assert "blank_or_na" in document_text


def test_template6_docx_structure_grouping_and_legends(tmp_path: Path) -> None:
    table = Template6ContentValidityTable(
        id="table.6",
        rows=(
            Template6ContentValidityRow(
                id="t6row.1",
                row_kind=Template6RowKind.BOX_SUMMARY,
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale=None,
                study_id="study.1",
                study_display_label="Hafner et al., 2022",
                cosmin_box="box_1_prom_development",
                measurement_property="prom_development",
                box_rating="doubtful",
                item_code=None,
                item_rating=None,
                uncertainty_status="reviewer_required",
                reviewer_decision_status="pending",
            ),
            Template6ContentValidityRow(
                id="t6row.2",
                row_kind=Template6RowKind.ITEM,
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale=None,
                study_id="study.1",
                study_display_label="Hafner et al., 2022",
                cosmin_box="box_1_prom_development",
                measurement_property="prom_development",
                box_rating=None,
                item_code="B1.1_target_population_definition",
                item_rating=None,
                uncertainty_status="reviewer_required",
                reviewer_decision_status="pending",
            ),
            Template6ContentValidityRow(
                id="t6row.3",
                row_kind=Template6RowKind.BOX_SUMMARY,
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale=None,
                study_id="study.1",
                study_display_label="Hafner et al., 2022",
                cosmin_box="box_2_content_validity",
                measurement_property="content_validity",
                box_rating="doubtful",
                item_code=None,
                item_rating=None,
                uncertainty_status="reviewer_required",
                reviewer_decision_status="pending",
            ),
        ),
        legends=(
            TableLegendEntry(key="box_summary", description="Box-level summary row."),
            TableLegendEntry(key="item", description="Item-level row."),
        ),
    )

    output_path = tmp_path / "template6.docx"
    export_template6_docx(table=table, output_path=output_path)

    doc = Document(str(output_path))
    assert len(doc.tables) == 1
    word_table = doc.tables[0]

    headers = [cell.text for cell in word_table.rows[0].cells]
    assert headers == [
        "PROM",
        "Version",
        "Subscale",
        "Study",
        "COSMIN Box",
        "Measurement Property",
        "Box Rating",
        "Item Code",
        "Item Rating",
        "Uncertainty Status",
        "Reviewer Status",
    ]
    assert _has_repeat_header(word_table.rows[0]) is True

    rows = [[cell.text for cell in row.cells] for row in word_table.rows[1:]]
    assert rows[0][0] == "PROM-X"
    assert rows[0][3] == "Hafner et al., 2022"
    assert rows[1][0] == ""
    assert rows[1][3] == ""
    assert rows[1][7] == "B1.1_target_population_definition"
    assert rows[1][8] == ""

    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Legend" in document_text
    assert "box_summary" in document_text
    assert "item" in document_text


def test_template7_docx_structure_and_summary_rows(tmp_path: Path) -> None:
    table = Template7EvidenceTable(
        id="table.7",
        rows=(
            Template7EvidenceRow(
                id="t7row.1",
                row_kind=Template7RowKind.STUDY,
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                measurement_property="reliability",
                study_id="study.1",
                study_display_label="Hafner et al., 2022",
                study_order_within_instrument_property=1,
                is_additional_study_row=False,
                per_study_rob="adequate",
                per_study_result="icc=0.81",
                study_rating="+",
            ),
            Template7EvidenceRow(
                id="t7row.2",
                row_kind=Template7RowKind.STUDY,
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                measurement_property="reliability",
                study_id="study.2",
                study_display_label="Hafner et al., 2022",
                study_order_within_instrument_property=2,
                is_additional_study_row=True,
                per_study_rob="doubtful",
                per_study_result="icc=0.59",
                study_rating="-",
            ),
            Template7EvidenceRow(
                id="t7row.3",
                row_kind=Template7RowKind.SUMMARY,
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                measurement_property="reliability",
                summarized_result="mixed findings across studies",
                overall_rating="±",
                certainty_of_evidence="low",
                total_sample_size=130,
            ),
            Template7EvidenceRow(
                id="t7row.4",
                row_kind=Template7RowKind.SUMMARY,
                instrument_name="PROM-X",
                instrument_version="v2",
                subscale="pain",
                measurement_property="reliability",
                summarized_result="stable sufficient evidence",
                overall_rating="+",
                certainty_of_evidence="moderate",
                total_sample_size=95,
            ),
        ),
        legends=(
            TableLegendEntry(key="±", description="Inconsistent findings."),
            TableLegendEntry(
                key="blank_or_na",
                description="Blank cells indicate NA or not assessed.",
            ),
        ),
    )

    output_path = tmp_path / "template7.docx"
    export_template7_docx(table=table, output_path=output_path)

    doc = Document(str(output_path))
    word_table = doc.tables[0]

    headers = [cell.text for cell in word_table.rows[0].cells]
    assert headers == [
        "PROM",
        "Version",
        "Subscale",
        "Measurement Property",
        "Study",
        "RoB",
        "Per-study Result",
        "Study Rating",
        "Summarized Result",
        "Overall Rating",
        "Certainty of Evidence",
        "Total n",
    ]
    assert _has_repeat_header(word_table.rows[0]) is True

    rows = [[cell.text for cell in row.cells] for row in word_table.rows[1:]]
    # Additional study row keeps grouping by blanking PROM/version/subscale/property.
    assert rows[1][:4] == ["", "", "", ""]
    assert rows[0][4] == "Hafner et al., 2022"
    assert rows[1][4] == "Hafner et al., 2022"
    # Summary row under same group should retain blank group columns with explicit label.
    assert rows[2][:4] == ["", "", "", ""]
    assert rows[2][4] == "Summary"
    assert rows[2][9] == "±"
    assert rows[2][10] == "low"
    assert rows[2][11] == "130"
    # New version summary should start a new visible group.
    assert rows[3][0] == "PROM-X"
    assert rows[3][1] == "v2"
    assert rows[3][4] == "Summary"
    assert word_table.rows[3].cells[11].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Legend" in document_text
    assert "blank_or_na" in document_text


def test_template8_docx_structure_grouping_and_legends(tmp_path: Path) -> None:
    table = Template8SummaryTable(
        id="table.8",
        rows=(
            Template8SummaryRow(
                id="t8row.1",
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                measurement_property="reliability",
                summarized_result="mixed",
                overall_rating="±",
                certainty_of_evidence="low",
                total_sample_size=130,
                inconsistent_findings=True,
            ),
            Template8SummaryRow(
                id="t8row.2",
                instrument_name="PROM-X",
                instrument_version="v1",
                subscale="pain",
                measurement_property="responsiveness",
                summarized_result=None,
                overall_rating=None,
                certainty_of_evidence=None,
                total_sample_size=None,
                inconsistent_findings=None,
            ),
            Template8SummaryRow(
                id="t8row.3",
                instrument_name="PROM-X",
                instrument_version="v2",
                subscale="pain",
                measurement_property="reliability",
                summarized_result="sufficient",
                overall_rating="+",
                certainty_of_evidence="moderate",
                total_sample_size=95,
                inconsistent_findings=False,
            ),
        ),
        legends=(
            TableLegendEntry(
                key="certainty_levels",
                description="high/moderate/low/very_low",
            ),
            TableLegendEntry(
                key="blank_or_na",
                description="Blank cells indicate NA or not assessed.",
            ),
        ),
    )

    output_path = tmp_path / "template8.docx"
    export_template8_docx(table=table, output_path=output_path)

    doc = Document(str(output_path))
    word_table = doc.tables[0]

    headers = [cell.text for cell in word_table.rows[0].cells]
    assert headers == [
        "PROM",
        "Version",
        "Subscale",
        "Measurement Property",
        "Summarized Result",
        "Overall Rating",
        "Certainty of Evidence",
        "Total n",
        "Inconsistent Findings",
    ]
    assert _has_repeat_header(word_table.rows[0]) is True

    rows = [[cell.text for cell in row.cells] for row in word_table.rows[1:]]
    # Same PROM/version/subscale on second row should blank grouping columns.
    assert rows[1][:3] == ["", "", ""]
    assert rows[1][3] == "responsiveness"
    assert rows[1][4] == ""
    # New version starts new group with visible identifiers.
    assert rows[2][0] == "PROM-X"
    assert rows[2][1] == "v2"
    assert rows[0][8] == "yes"
    assert rows[2][8] == "no"
    assert word_table.rows[1].cells[7].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Legend" in document_text
    assert "certainty_levels" in document_text


def _has_repeat_header(row: _Row) -> bool:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return False
    return tr_pr.find(qn("w:tblHeader")) is not None


def _cell_fill(cell: _Cell) -> str | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))
